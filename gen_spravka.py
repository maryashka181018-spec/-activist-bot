"""
Генерация справки-подтверждения на основе оригинального шаблона.
Копирует первую страницу оригинала, меняет нужные поля.
"""
import copy, io, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

TEMPLATE_PATH = "СПРАВКА_ПОДТВЕРЖДЕНИЕ.docx"

def find_and_replace_in_paragraph(p, old_text, new_text):
    """Заменяет текст в параграфе сохраняя форматирование первого run'а"""
    full = "".join(r.text for r in p.runs)
    if old_text not in full:
        return False
    # Очищаем все runs кроме первого, первому ставим новый текст
    new_full = full.replace(old_text, new_text)
    if p.runs:
        p.runs[0].text = new_full
        for r in p.runs[1:]:
            r.text = ""
    return True

def clone_paragraph_after(ref_para, template_para):
    """Вставляет клон параграфа после ref_para"""
    new_p = copy.deepcopy(template_para._element)
    ref_para._element.addnext(new_p)
    return ref_para.__class__(new_p, ref_para._parent)

def generate_spravka(events_data, template_path=TEMPLATE_PATH):
    """
    events_data: список dict с ключами title, date, location, students (list of str)
    Возвращает bytes готового docx
    """
    # Загружаем шаблон — первая справка уже там
    doc = Document(template_path)
    paras = doc.paragraphs

    # Находим индексы ключевых параграфов первой справки
    intro_idx = None      # "В рамках мероприятия..."
    first_student_idx = None  # первый студент
    last_student_idx = None   # последний студент перед "в указанный период"
    after_students_idx = None # "в указанный период..."

    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("В рамках мероприятия") and intro_idx is None:
            intro_idx = i
        if intro_idx and first_student_idx is None and i > intro_idx:
            if t and not t.startswith("в указанный"):
                first_student_idx = i
        if first_student_idx and not after_students_idx:
            if t.startswith("в указанный"):
                after_students_idx = i
                last_student_idx = i - 1

    # Шаблон студента — берём форматирование из первого студента оригинала
    student_template_para = paras[first_student_idx]

    # ── Правим первую справку (events_data[0]) ──────────────────────────────
    ev0 = events_data[0]

    # Меняем строку "проведенного X в Y"
    intro_p = paras[intro_idx]
    old_intro = intro_p.text
    # Вычленяем шаблонную часть — заменяем дату и место
    new_intro = re.sub(
        r'проведенного .+ в .+ обучающиеся:',
        f"проведенного {ev0['date']} в {ev0['location']} обучающиеся:",
        old_intro
    )
    if intro_p.runs:
        full = "".join(r.text for r in intro_p.runs)
        intro_p.runs[0].text = new_intro
        for r in intro_p.runs[1:]:
            r.text = ""

    # Удаляем старых студентов
    old_student_paras = paras[first_student_idx:after_students_idx]
    for sp in old_student_paras:
        sp._element.getparent().remove(sp._element)

    # Вставляем новых студентов перед "в указанный период"
    after_p = doc.paragraphs[intro_idx + 1]  # после правки это "в указанный"
    ref = intro_p
    for student in reversed(ev0["students"]):
        new_p = copy.deepcopy(student_template_para._element)
        # Меняем текст
        for r in new_p.findall('.//' + qn('w:r')):
            t_el = r.find(qn('w:t'))
            if t_el is not None:
                t_el.text = student
                break
        ref._element.addnext(new_p)

    # ── Добавляем остальные справки (events_data[1:]) ──────────────────────
    # Находим разрыв страницы — клонируем всё до него
    page_break_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'pageBreak' in p._element.xml:
            page_break_idx = i
            break

    template_block = doc.paragraphs[:page_break_idx] if page_break_idx else doc.paragraphs[:]

    for ev in events_data[1:]:
        # Добавляем разрыв страницы
        pb = OxmlElement('w:p')
        r_el = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r_el.append(br)
        pb.append(r_el)
        doc.element.body.append(pb)

        # Клонируем все параграфы шаблона
        cloned = []
        for tp in template_block:
            new_el = copy.deepcopy(tp._element)
            doc.element.body.append(new_el)
            cloned.append(new_el)

        # Правим клон — ищем "В рамках" и студентов
        c_paras = [doc.paragraphs[-(len(cloned) - i)] for i in range(len(cloned))]

        c_intro = None
        c_first_st = None
        c_after_st = None
        for cp in c_paras:
            t = cp.text.strip()
            if t.startswith("В рамках") and c_intro is None:
                c_intro = cp
            elif c_intro and c_first_st is None and t and not t.startswith("в указанный"):
                c_first_st = cp
            elif c_first_st and c_after_st is None and t.startswith("в указанный"):
                c_after_st = cp

        if c_intro:
            old_t = c_intro.text
            new_t = re.sub(
                r'проведенного .+ в .+ обучающиеся:',
                f"проведенного {ev['date']} в {ev['location']} обучающиеся:",
                old_t
            )
            if c_intro.runs:
                c_intro.runs[0].text = new_t
                for r in c_intro.runs[1:]:
                    r.text = ""

        # Удаляем старых студентов клона
        if c_first_st and c_after_st:
            curr = c_first_st._element
            stop = c_after_st._element
            to_remove = []
            while curr != stop:
                to_remove.append(curr)
                curr = curr.getnext()
            for el in to_remove:
                el.getparent().remove(el)

            # Вставляем новых
            for student in reversed(ev["students"]):
                new_p = copy.deepcopy(student_template_para._element)
                for r in new_p.findall('.//' + qn('w:r')):
                    t_el = r.find(qn('w:t'))
                    if t_el is not None:
                        t_el.text = student
                        break
                c_intro._element.addnext(new_p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

